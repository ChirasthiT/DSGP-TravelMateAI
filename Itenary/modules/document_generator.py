from docx import Document
import os

def generate_word_document(itinerary_text, template_path, arrival_date, departure_date):
    """Generate a Word document using the template and itinerary text"""
    # Load template
    doc = Document("Itenary/resources/template.docx")
    
    # Add title
    doc.add_heading(f'Travel Itinerary ({arrival_date} to {departure_date})', 0)
    
    # Add itinerary content
    doc.add_paragraph(itinerary_text)
    
    # Save document
    output_path = 'temp_itinerary.docx'
    doc.save(output_path)
    return output_path