from flask import Blueprint, render_template, request, send_file, jsonify
from Itenary.modules.keyword_extraction import process_travel_description
from Itenary.modules.recommender import create_integrated_recommender
from Itenary.modules.itinerary_generator import ItineraryGenerator
from Itenary.modules.utils import format_extracted_info
from datetime import datetime
import pandas as pd
import os
from docx import Document
from io import BytesIO
from docx.shared import Inches
from dotenv import load_dotenv
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from Itenary.modules.constants import VALID_LOCATIONS

load_dotenv()

# Create a Blueprint for the Itinerary component
itinerary_blueprint = Blueprint('itinerary', __name__, template_folder='templates', static_folder='static')

# Helper function to load activities
def load_location_activities():
    df = pd.read_csv('./data/Location_and_Activities (FINAL).csv', encoding='latin-1')
    location_activities = {}
    for _, row in df.iterrows():
        activities = [a.strip() for a in row['Activity'].split(',')]
        location_activities[row['Location'].lower()] = activities
    return location_activities

@itinerary_blueprint.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            # Store form data
            user_data = {
                'arrival_date': request.form['arrival_date'],
                'arrival_time': request.form['arrival_time'],
                'departure_date': request.form['departure_date'],
                'departure_time': request.form['departure_time'],
                'description': request.form['description']
            }
            
            # Process initial recommendations
            extracted_info_str = process_travel_description(user_data['description'])
            extracted_info = format_extracted_info(extracted_info_str)
            recommendations = create_integrated_recommender(extracted_info_str)
            
            # Get all activities for locations
            location_activities = load_location_activities()
            locations_activities = {}
            for loc in extracted_info['locations']:
                loc_lower = loc.lower()
                if loc_lower in location_activities:
                    recommended = set(recommendations.get(loc, []))
                    all_activities = location_activities[loc_lower]
                    activities = [{
                        'name': a,
                        'checked': a in recommended
                    } for a in all_activities]
                    locations_activities[loc] = activities
            
            return render_template('adjust.html', 
                                 locations_activities=locations_activities,
                                 user_data=user_data)
                                 
        except Exception as e:
            print(f"Error processing request: {str(e)}")
            error_message = "An error occurred while processing your request. Please try again."
            return render_template('index.html', 
                                 error=error_message, 
                                 locations=sorted([loc.title() for loc in VALID_LOCATIONS]))
    
    # GET request
    return render_template('index.html', 
                         locations=sorted([loc.title() for loc in VALID_LOCATIONS]))

@itinerary_blueprint.route('/generate', methods=['POST'])
def generate():
    try:
        # Collect selected activities
        selected_activities = {}
        for key in request.form:
            if key.startswith('activity_'):
                try:
                    parts = key[len('activity_'):].split('|', 1)
                    if len(parts) != 2:
                        continue
                    loc, activity = parts
                    loc = loc.lower().strip()
                    if loc not in selected_activities:
                        selected_activities[loc] = []
                    selected_activities[loc].append(activity.strip())
                except Exception as e:
                    print(f"Error processing activity key {key}: {str(e)}")
                    continue
        
        # Get form data
        user_data = {
            'arrival_date': request.form['arrival_date'],
            'arrival_time': request.form['arrival_time'],
            'departure_date': request.form['departure_date'],
            'departure_time': request.form['departure_time'],
            'description': request.form['description']
        }
        
        # Generate itinerary
        extracted_info_str = process_travel_description(user_data['description'])
        recommendations = create_integrated_recommender(extracted_info_str)
        
        # Update recommendations with user selections
        for loc in recommendations:
            if loc.lower() in selected_activities:
                recommendations[loc] = selected_activities[loc.lower()]
        
         # Generate with updated prompt structure
        generator = ItineraryGenerator(os.getenv('GROQ_API_KEY'))
        formatted_info = format_extracted_info(extracted_info_str)
        prompt = generator.create_prompt(
            formatted_info,
            recommendations,
            user_data['arrival_date'],
            user_data['arrival_time'],
            user_data['departure_date'],
            user_data['departure_time']
        )
        
        itinerary_data = generator.generate_itinerary(prompt)
        
        # Create Word document with tables
        doc = Document('resources/template.docx')
        
        # Set default font to Calibri for the entire document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Add heading with Calibri font
        heading = doc.add_heading('Travel Itinerary', 0)
        for run in heading.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(24)

        # Main itinerary table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Normal Table'
        hdr = table.rows[0].cells
        
        # Set bold header text with Calibri font
        headers = ['Location', 'Check-In', 'Check-Out', 'Nights', 'Daily Itinerary']
        for i, header in enumerate(headers):
            hdr[i].text = header
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'

        # Set column widths
        widths = [2.5, 1.0, 1.0, 0.5, 18.0]  # In inches
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
                # Set Calibri font for all cells
                for paragraph in row.cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'

        # Populate table rows
        for entry in itinerary_data.get('itinerary', []):
            row = table.add_row().cells
            row[0].text = entry.get('location', '')
            row[1].text = entry.get('check_in', '')
            row[2].text = entry.get('check_out', '')
            row[3].text = str(entry.get('nights', ''))
            row[4].text = entry.get('details', '')
            
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'

        # Add fees table if available
        if itinerary_data.get('fees'):
            doc.add_paragraph()  # Add spacing
            heading = doc.add_heading('Activity Costs', level=1)
            for run in heading.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(24)

            fees_table = doc.add_table(rows=1, cols=3)
            fees_table.style = 'Normal Table'
            f_hdr = fees_table.rows[0].cells
            
            # Set bold header text with Calibri font
            fee_headers = ['Location', 'Activity', 'Approximate Rate (USD)']
            for i, header in enumerate(fee_headers):
                f_hdr[i].text = header
                for paragraph in f_hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = 'Calibri'
                    
            for fee in itinerary_data['fees']:
                row = fees_table.add_row().cells
                row[0].text = fee.get('location', '')
                row[1].text = fee.get('activity', '')
                row[2].text = f"${fee.get('rate', '')}"
                    
                for cell in row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                    
        # Add additional content from another document
        additional_doc = Document('./resources/additional_details.docx')
        for element in additional_doc.element.body:
            doc.element.body.append(element)   

        # Save to bytes buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"itinerary_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        
        # Return both the file and a status indicator
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        print(f"Error generating itinerary: {str(e)}")
        return jsonify({'error': 'Failed to generate itinerary'}), 500