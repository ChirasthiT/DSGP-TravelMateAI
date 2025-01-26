import os
from dotenv import load_dotenv
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from datetime import datetime, timedelta
import groq
import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Initialize Groq client
client = groq.Groq(api_key=GROQ_API_KEY)

# Constants
DEFAULT_PASSWORD = "admin2025"
LOCATIONS = ["Kandy", "Sigiriya", "Habarana", "Dambulla", "Nuwara Eliya", "Ella",
             "Hatton", "Bentota", "Galle", "Yala", "Colombo", "Anuradhapura",
             "Polonnaruwa", "Trincomalee", "Jaffna", "Pasikuda", "Arugam Bay"]




def create_llm_prompt(locations, travel_companion, arrival_date, arrival_time,
                      departure_date, departure_time, additional_details):
    days = (departure_date - arrival_date).days + 1

    prompt = f"""You are a tourism company. Create a detailed Sri Lanka travel itinerary based on these specifications for the customers:

Travel Details:
- Locations: {', '.join(locations)}
- Traveling with: {travel_companion}
- Duration: {days} days
- Arrival: {arrival_date.strftime('%Y-%m-%d')} ({arrival_time})
- Departure: {departure_date.strftime('%Y-%m-%d')} ({departure_time})
- Additional Notes: {additional_details if additional_details else 'None'}

Create a logical day-by-day itinerary considering travel times between locations.
For each location, include specific activities and entrance fees.
The customer will be arriving from the airport and departing from the airport at the end
Whenever applicable start with 'Welcome at the airport and transfer to a hotel near airport for a short stay.'

Return only a valid JSON object in this exact structure:
{{
    "itinerary": [
        {{
            "location": "Location Name (Approx: travel time Xhrs from previous location)",
            "check_in": "YYYY-MM-DD",
            "check_out": "YYYY-MM-DD",
            "nights": X,
            "details": "Detailed day-by-day activities [ex: Day 1: Welcome at the airport and transfer to a hotel near airport for a short stay. Visit Pinnawala elephant orphanage en-route Sigiriya. Engage in elephant walking, bathing & feeding at Millennium Foundation. Reach Sigiriya in the afternoon and engage in a Jeep safari in the wilderness.  \\nDay 2: After breakfast, climb Sigiriya rock fortress. Climb Dambulla cave temple & visit Golden temple. Engage in a Village tour – experience authentic village life including visiting a farm, climbing a tree house, catamaran ride through village’s lake and help prepare lunch.]"
        }}
    ],
    "fees": [
        {{
            "location": "Location Name",
            "activity": "Activity Name",
            "rate": XX
        }}
    ]
}}
(NO PREAMBLE)"""

    return prompt


def generate_llm_response(prompt):
    try:
        completion = client.chat.completions.create(
            model="llama-3.1-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )
        response_text = completion.choices[0].message.content.strip()

        # Extract JSON content if wrapped in code blocks
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].strip()

        # Parse JSON and validate structure
        data = json.loads(response_text)
        if not isinstance(data, dict) or 'itinerary' not in data or 'fees' not in data:
            raise ValueError("Invalid response structure")

        return data
    except json.JSONDecodeError as e:
        st.error(f"Error parsing response: {str(e)}")
        return None
    except Exception as e:
        st.error(f"Error generating itinerary: {str(e)}")
        return None


def set_cell_border(cell, border_style='single', border_width=4):
    """Set cell borders with specified style and width"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create border elements
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), border_style)
        border.set(qn('w:sz'), str(border_width))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)


def apply_table_style(table):
    """Apply consistent styling to a table with strong borders"""
    for row in table.rows:
        for cell in row.cells:
            # Set strong borders
            set_cell_border(cell)

            # Apply font styling to all paragraphs in the cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(14)


def generate_itinerary_document(data):
    try:
        doc = Document("App/Assets/template.docx")

        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(14)

    except Exception as e:
        st.error(f"Error loading template: {str(e)}")
        return None

    # Add headings with Calibri font
    heading = doc.add_heading("Travel Itinerary", level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'

    # Create main itinerary table
    table = doc.add_table(rows=1, cols=5)

    widths = [1.5, 1.2, 1.2, 1.2, 3.9]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])

    # Add and style headers
    headers = ["Location", "Check-in", "Check-Out", "Number of Nights", "Itinerary"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.font.bold = True

    # Add and style data rows
    for item in data['itinerary']:
        row = table.add_row()
        row.cells[0].text = item['location']
        row.cells[1].text = item['check_in']
        row.cells[2].text = item['check_out']
        row.cells[3].text = str(item['nights'])
        row.cells[4].text = item['details']

        for cell in row.cells:
            set_cell_border(cell)

    apply_table_style(table)

    doc.add_paragraph()
    heading = doc.add_heading("Entrance & Activity Fees", level=1)
    for run in heading.runs:
        run.font.name = 'Calibri'

    # Create fees table
    fees_table = doc.add_table(rows=1, cols=3)

    # Add and style headers
    headers = ["Location", "Attraction/Activity", "Approximate Rate per Adult (USD)"]
    for i, header in enumerate(headers):
        cell = fees_table.rows[0].cells[i]
        cell.text = header
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(14)
                run.font.bold = True

    # Add and style data rows
    for fee in data['fees']:
        row = fees_table.add_row()
        row.cells[0].text = fee['location']
        row.cells[1].text = fee['activity']
        row.cells[2].text = f"${fee['rate']}"

        for cell in row.cells:
            set_cell_border(cell)

    apply_table_style(fees_table)

    try:
        output_path = "Generated_Itinerary.docx"
        doc.save(output_path)
        return output_path
    except Exception as e:
        st.error(f"Error saving document: {str(e)}")
        return None




def main():
    st.title("Heart Attached Holidays - Itinerary Builder")

    if 'logged_in' not in st.session_state:
        st.session_state['logged_in'] = False
    if 'admin_password' not in st.session_state:
        st.session_state['admin_password'] = DEFAULT_PASSWORD
    if 'itinerary_generated' not in st.session_state:
        st.session_state['itinerary_generated'] = False
    if 'output_file' not in st.session_state:
        st.session_state['output_file'] = None

    if not st.session_state['logged_in']:
        with st.form("login_form"):
            password = st.text_input("Enter password:", type="password")
            submitted = st.form_submit_button("Login")
            if submitted:
                if password == st.session_state['admin_password']:
                    st.session_state['logged_in'] = True
                    st.success("Login successful!")
                    st.rerun()
                else:
                    st.error("Incorrect password.")
    else:
        with st.sidebar:
            st.subheader("Admin Options")

            with st.expander("Update Password"):
                current_pwd = st.text_input("Current password:", type="password")
                new_pwd = st.text_input("New password:", type="password")
                confirm_pwd = st.text_input("Confirm new password:", type="password")
                if st.button("Update Password"):
                    if current_pwd == st.session_state['admin_password']:
                        if new_pwd == confirm_pwd:
                            st.session_state['admin_password'] = new_pwd
                            st.success("Password updated successfully!")
                        else:
                            st.error("New passwords don't match!")
                    else:
                        st.error("Current password is incorrect!")

            if st.button("Logout"):
                st.session_state['logged_in'] = False
                st.rerun()

        st.subheader("Create Itinerary")
        st.markdown("Fields marked with * are required")

        # Form for itinerary details
        with st.form("itinerary_form"):
            locations = st.multiselect("Select Locations *",
                                     LOCATIONS,
                                     help="Select one or more locations",
                                     key="locations")

            additional_locs = st.text_input("Additional Locations (Optional):",
                                          help="Enter additional locations separated by commas")

            if additional_locs:
                locations.extend([loc.strip() for loc in additional_locs.split(',')])

            travel_companion = st.selectbox(
                "With whom are you travelling *",
                ["Family", "Friends", "Partner", "Solo"],
                help="Select your travel companion type",
                key="travel_companion"
            )

            col1, col2 = st.columns(2)
            with col1:
                arrival_date = st.date_input(
                    "Date of Arrival *",
                    min_value=datetime.today(),
                    help="Select your arrival date",
                    key="arrival_date"
                )
                arrival_time = st.selectbox(
                    "Time of Arrival *",
                    ["Morning", "Noon", "Evening", "Night"],
                    key="arrival_time"
                )

            with col2:
                departure_date = st.date_input(
                    "Date of Departure *",
                    min_value=arrival_date,
                    help="Select your departure date",
                    key="departure_date"
                )
                departure_time = st.selectbox(
                    "Time of Departure *",
                    ["Morning", "Noon", "Evening", "Night"],
                    key="departure_time"
                )

            additional_details = st.text_area(
                "Additional Details (Optional):",
                max_chars=500,
                help="Enter any additional requirements or preferences (max 500 characters)"
            )

            generate_button = st.form_submit_button("Generate Itinerary")

            if generate_button:
                if not locations:
                    st.error("Please select at least one location.")
                elif arrival_date >= departure_date:
                    st.error("Departure date must be after arrival date.")
                else:
                    with st.spinner("Generating itinerary... This may take a few moments."):
                        try:
                            prompt = create_llm_prompt(
                                locations, travel_companion, arrival_date, arrival_time,
                                departure_date, departure_time, additional_details
                            )

                            itinerary_data = generate_llm_response(prompt)

                            if itinerary_data:
                                st.session_state['output_file'] = generate_itinerary_document(itinerary_data)
                                if st.session_state['output_file']:
                                    st.session_state['itinerary_generated'] = True
                                    st.rerun()
                        except Exception as e:
                            st.error(f"An error occurred: {str(e)}")

        # Download button outside the form
        if st.session_state['itinerary_generated'] and st.session_state['output_file']:
            st.success("Itinerary generated successfully!")
            with open(st.session_state['output_file'], "rb") as file:
                st.download_button(
                    "Download Itinerary",
                    data=file,
                    file_name="Sri_Lanka_Itinerary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )



if __name__ == "__main__":
    main()